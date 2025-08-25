import streamlit as st
import pandas as pd
import numpy as np
from docx import Document
from io import BytesIO
from scipy import stats

def show_table1_tab():
    st.header("Table 1 生成器")
    table_file = st.file_uploader("上传 GBD CSV 文件用于生成 Table 1", type=["csv"], key="table")

    if table_file is not None:
        df = pd.read_csv(table_file)

        st.subheader("选择生成参数")
        available_causes = df['cause'].dropna().unique().tolist()
        available_ages = df['age'].dropna().unique().tolist()
        available_locations = df['location'].dropna().unique().tolist()
        available_sexes = df['sex'].dropna().unique().tolist()
        available_measures = df['measure'].dropna().unique().tolist()

        selected_cause = st.selectbox("请选择病因 (cause)", available_causes)
        selected_age = st.selectbox("请选择年龄组 (age)", available_ages)
        selected_sex = st.selectbox("请选择性别 (sex)", available_sexes)
        selected_measure = st.selectbox("请选择指标 (measure)", available_measures)

        year_range = [1990, 2021]

        def format_val(val, lower, upper, digits=2):
            return f"{round(val, digits)} ({round(lower, digits)} to {round(upper, digits)})"

        def get_row(location):
            row = {"Feature": location}

            for year in year_range:
                d_number = df[(df['location'] == location) & (df['year'] == year) & (df['cause'] == selected_cause) &
                              (df['metric'] == 'Number') & (df['measure'] == selected_measure) &
                              (df['age'] == ('All ages' if selected_age == 'Age-standardized' else selected_age)) &
                              (df['sex'] == selected_sex)]

                d_rate = df[(df['location'] == location) & (df['year'] == year) & (df['cause'] == selected_cause) &
                            (df['metric'] == 'Rate') & (df['measure'] == selected_measure) &
                            (df['age'] == selected_age) & (df['sex'] == selected_sex)]

                row[f"Cases_{year}"] = format_val(d_number['val'].sum(), d_number['lower'].sum(), d_number['upper'].sum(), 0) if not d_number.empty else "NA"
                row[f"Rates_{year}"] = format_val(d_rate['val'].sum(), d_rate['lower'].sum(), d_rate['upper'].sum()) if not d_rate.empty else "NA"

            d_1990 = df[(df['location'] == location) & (df['year'] == 1990) & (df['cause'] == selected_cause) &
                        (df['metric'] == 'Number') & (df['measure'] == selected_measure) &
                        (df['age'] == ('All ages' if selected_age == 'Age-standardized' else selected_age)) &
                        (df['sex'] == selected_sex)]
            d_2021 = df[(df['location'] == location) & (df['year'] == 2021) & (df['cause'] == selected_cause) &
                        (df['metric'] == 'Number') & (df['measure'] == selected_measure) &
                        (df['age'] == ('All ages' if selected_age == 'Age-standardized' else selected_age)) &
                        (df['sex'] == selected_sex)]
            if not d_1990.empty and not d_2021.empty:
                v1 = d_1990['val'].sum()
                v2 = d_2021['val'].sum()
                change = ((v2 - v1) / v1) * 100 if v1 != 0 else np.nan
                row['Cases_change'] = f"{round(change, 2)}%"
            else:
                row['Cases_change'] = "NA"

            d_rate_all = df[(df['location'] == location) & (df['cause'] == selected_cause) &
                            (df['metric'] == 'Rate') & (df['measure'] == selected_measure) &
                            (df['age'] == selected_age) & (df['sex'] == selected_sex)]

            eapc_ci = "NA"
            if d_rate_all.shape[0] >= 2:
                try:
                    d_rate_all = d_rate_all[d_rate_all['val'] > 0]
                    d_rate_all['y'] = np.log(d_rate_all['val'])
                    slope, intercept, r_value, p_value, std_err = stats.linregress(d_rate_all['year'], d_rate_all['y'])
                    eapc = (np.exp(slope) - 1) * 100
                    lci = (np.exp(slope - 1.96 * std_err) - 1) * 100
                    uci = (np.exp(slope + 1.96 * std_err) - 1) * 100
                    eapc_ci = f"{round(eapc, 2)} ({round(lci, 2)} to {round(uci, 2)})"
                except:
                    pass
            row['EAPC_CI'] = eapc_ci

            return row

        if st.button("生成 Table 1"):
            summary_df = []
            all_locations = sorted(df['location'].dropna().unique().tolist())

            for loc in all_locations:
                summary_df.append(get_row(loc))

            summary_table = pd.DataFrame(summary_df)
            st.dataframe(summary_table)

            doc = Document()
            doc.add_heading("Table 1", level=1)

            table = doc.add_table(rows=1, cols=len(summary_table.columns))
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(summary_table.columns):
                hdr_cells[i].text = str(col)

            for index, row in summary_table.iterrows():
                row_cells = table.add_row().cells
                for i, item in enumerate(row):
                    row_cells[i].text = str(item)

            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)

            st.download_button(
                label="下载 Word 文件",
                data=buf,
                file_name="Table1.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )