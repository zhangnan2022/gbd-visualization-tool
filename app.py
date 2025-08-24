# ⚠️ 这是合并并优化后的 Streamlit 应用代码：实现分组趋势图 + Table 1 多选输出，按 GBD 官网风格美化，作者名展示，支持 Word 下载

import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import altair as alt
import io
import base64
from docx import Document
from docx.shared import Inches
from io import BytesIO
import math
from scipy import stats

st.set_page_config(page_title="GBD 可视化助手", layout="wide")
st.title("GBD 可视化助手 · 内测版")
st.markdown("作者：zhangnan")

# 读取排序 order.csv（嵌入代码中，不用用户上传）
ORDER = []
with open("/mnt/data/order.csv", "r") as f:
    ORDER = [line.strip() for line in f.readlines() if line.strip()]

# 用于 Table 1 分组标题
SDI_GROUPS = ['High SDI', 'High-middle SDI', 'Middle SDI', 'Low-middle SDI', 'Low SDI']

# 页面结构
tabs = st.tabs(["分组趋势图", "Table 1 生成器"])

# 分组趋势图
with tabs[0]:
    st.header("分组趋势图")
    uploaded_file = st.file_uploader("上传 GBD CSV 数据文件", type=["csv"], key="chart")

    if uploaded_file is not None:
        df = pd.read_csv(uploaded_file)

        # 获取选项
        measure_options = df['measure'].dropna().unique().tolist()
        location_options = df['location'].dropna().unique().tolist()
        sex_options = df['sex'].dropna().unique().tolist()
        age_options = df['age'].dropna().unique().tolist()
        cause_options = df['cause'].dropna().unique().tolist()
        metric_options = df['metric'].dropna().unique().tolist()
        year_options = sorted(df['year'].dropna().unique().tolist())

        st.subheader("筛选条件")
        chart_type = st.selectbox("选择图表类型", ["折线图"], index=0)
        group_by = st.selectbox("选择分组字段", ["sex", "age", "location", "cause"], index=0)

        selected_measure = st.selectbox("请选择 measure", ["请选择"] + measure_options)
        selected_location = st.selectbox("请选择 location", ["请选择"] + location_options)
        selected_sex = st.multiselect("请选择 sex", sex_options) if group_by != "sex" else []
        selected_age = st.multiselect("请选择 age", age_options) if group_by != "age" else []
        selected_cause = st.selectbox("请选择 cause", ["请选择"] + cause_options)
        selected_metric = st.selectbox("请选择 metric", ["请选择"] + metric_options)
        selected_years = st.slider("选择年份范围", min_value=min(year_options), max_value=max(year_options), value=(min(year_options), max(year_options)))

        show_ci = st.checkbox("显示置信区间 (upper/lower)")

        if selected_measure != "请选择" and selected_location != "请选择" and selected_cause != "请选择" and selected_metric != "请选择":
            chart_df = df[(df['measure'] == selected_measure) &
                          (df['location'] == selected_location) &
                          (df['cause'] == selected_cause) &
                          (df['metric'] == selected_metric) &
                          (df['year'] >= selected_years[0]) & (df['year'] <= selected_years[1])]

            if group_by != "sex" and selected_sex:
                chart_df = chart_df[chart_df['sex'].isin(selected_sex)]
            if group_by != "age" and selected_age:
                chart_df = chart_df[chart_df['age'].isin(selected_age)]

            if not chart_df.empty:
                base = alt.Chart(chart_df).encode(
                    x=alt.X('year:O', title='Year'),
                    y=alt.Y('val:Q', title='Value'),
                    color=group_by + ":N"
                )
                lines = base.mark_line(point=True)
                if show_ci:
                    band = base.mark_area(opacity=0.3).encode(y='lower:Q', y2='upper:Q')
                    chart = (band + lines).properties(width=800, height=400)
                else:
                    chart = lines.properties(width=800, height=400)
                st.altair_chart(chart, use_container_width=True)

# Table 1 生成器
with tabs[1]:
    st.header("Table 1 生成器")
    table_file = st.file_uploader("上传 GBD CSV 文件用于生成 Table 1", type=["csv"], key="table")

    if table_file is not None:
        df = pd.read_csv(table_file)

        available_causes = df['cause'].dropna().unique().tolist()
        available_ages = df['age'].dropna().unique().tolist()
        available_locations = df['location'].dropna().unique().tolist()
        available_sexes = df['sex'].dropna().unique().tolist()

        st.subheader("请选择纳入 Table 1 的组合")
        selected_cause = st.selectbox("病因 (cause)", available_causes)
        selected_ages = st.multiselect("年龄组 (age)", available_ages)
        selected_sexes = st.multiselect("性别 (sex)", available_sexes)
        selected_locations = st.multiselect("地区 (location)", available_locations)

        def format_val(val, lower, upper, digits=2):
            return f"{round(val, digits)} ({round(lower, digits)} to {round(upper, digits)})"

        def get_row(loc, sex, age):
            row = {"Feature": f"{loc} ({sex}, {age})"}
            for year in [1990, 2021]:
                d_num = df[(df['location'] == loc) & (df['year'] == year) & (df['cause'] == selected_cause) &
                           (df['metric'] == 'Number') & (df['measure'] == 'Prevalence') &
                           (df['sex'] == sex) & (df['age'] == age)]
                d_rate = df[(df['location'] == loc) & (df['year'] == year) & (df['cause'] == selected_cause) &
                            (df['metric'] == 'Rate') & (df['measure'] == 'Prevalence') &
                            (df['sex'] == sex) & (df['age'] == age)]
                if not d_num.empty:
                    row[f"Cases_{year}"] = format_val(d_num['val'].sum(), d_num['lower'].sum(), d_num['upper'].sum(), 0)
                else:
                    row[f"Cases_{year}"] = "NA"
                if not d_rate.empty:
                    row[f"Rates_{year}"] = format_val(d_rate['val'].sum(), d_rate['lower'].sum(), d_rate['upper'].sum())
                else:
                    row[f"Rates_{year}"] = "NA"

            # Change
            try:
                v1990 = df[(df['location'] == loc) & (df['year'] == 1990) & (df['cause'] == selected_cause) &
                           (df['metric'] == 'Number') & (df['measure'] == 'Prevalence') &
                           (df['sex'] == sex) & (df['age'] == age)]['val'].sum()
                v2021 = df[(df['location'] == loc) & (df['year'] == 2021) & (df['cause'] == selected_cause) &
                           (df['metric'] == 'Number') & (df['measure'] == 'Prevalence') &
                           (df['sex'] == sex) & (df['age'] == age)]['val'].sum()
                change = ((v2021 - v1990) / v1990) * 100 if v1990 else np.nan
                row['Cases_change'] = f"{round(change, 2)}%"
            except:
                row['Cases_change'] = "NA"

            # EAPC
            d_eapc = df[(df['location'] == loc) & (df['cause'] == selected_cause) &
                        (df['metric'] == 'Rate') & (df['measure'] == 'Prevalence') &
                        (df['sex'] == sex) & (df['age'] == age)]
            eapc_ci = "NA"
            if d_eapc.shape[0] >= 2:
                try:
                    d_eapc = d_eapc[d_eapc['val'] > 0]
                    d_eapc['y'] = np.log(d_eapc['val'])
                    slope, _, _, _, std_err = stats.linregress(d_eapc['year'], d_eapc['y'])
                    eapc = (np.exp(slope) - 1) * 100
                    lci = (np.exp(slope - 1.96 * std_err) - 1) * 100
                    uci = (np.exp(slope + 1.96 * std_err) - 1) * 100
                    eapc_ci = f"{round(eapc,2)} ({round(lci,2)} to {round(uci,2)})"
                except:
                    pass
            row['EAPC_CI'] = eapc_ci
            return row

        if st.button("生成 Table 1"):
            results = []
            group_title_set = set()
            for group_title in ["Global", "Sex group", "SDI regions", "Geographical regions"]:
                results.append({"Feature": group_title, "Cases_1990": "", "Rates_1990": "", "Cases_2021": "", "Rates_2021": "", "Cases_change": "", "EAPC_CI": ""})
                group_title_set.add(group_title)
                for sex in selected_sexes:
                    for age in selected_ages:
                        for loc in selected_locations:
                            if group_title == "Global" and loc == "Global":
                                results.append(get_row(loc, sex, age))
                            elif group_title == "Sex group" and loc == "Global":
                                results.append(get_row(loc, sex, age))
                            elif group_title == "SDI regions" and loc in SDI_GROUPS:
                                results.append(get_row(loc, sex, age))
                            elif group_title == "Geographical regions" and loc in ORDER and loc not in SDI_GROUPS and loc != "Global":
                                results.append(get_row(loc, sex, age))

            summary_df = pd.DataFrame(results)
            st.dataframe(summary_df)

            # 导出 Word
            doc = Document()
            doc.add_heading("Table 1", level=1)
            table = doc.add_table(rows=1, cols=len(summary_df.columns))
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(summary_df.columns):
                hdr_cells[i].text = str(col)
            for _, row in summary_df.iterrows():
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)
            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            st.download_button("下载 Table1.docx", data=buf, file_name="Table1.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
